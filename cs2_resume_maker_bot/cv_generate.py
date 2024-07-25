import telebot
from telebot.types import InputMediaPhoto
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import requests
from io import BytesIO
import os

# Replace 'YOUR_TELEGRAM_BOT_TOKEN_HERE' with your actual bot token
TOKEN = "6837439077:AAGaEwhRZCEuN6UGgIF8KKQkNG-DxemOvJg"
bot = telebot.TeleBot(TOKEN)

# Dictionary to store user data temporarily
user_data = {}

# Class to store user data
class UserData:
    def __init__(self):
        self.name = None
        self.surname = None
        self.age = None
        self.address = None
        self.objective = None
        self.education = None
        self.specialization = None
        self.skills = None
        self.languages = None
        self.photo_id = None
        self.experience = ""
# Function to create the modern CV
def create_modern_cv(chat_id):
    try:
        user = user_data[chat_id]
        doc = Document()
        setup_document_style(doc)

        # Adding CV title
        header = doc.add_heading(level=0)
        run = header.add_run('RESUME')
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 255, 0)  #light green
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()  # Adding extra space

        # Adding photo if available
        if user.photo_id:
            add_photo(doc, user.photo_id)

        # Adding objective section
        if user.objective:
            add_objective_section(doc, user.objective)

        # Adding personal information section
        add_info_section(doc, 'Personal Informations', [
            f"Name: {user.name} {user.surname}",
            f"Age: {user.age}",
            f"Address: {user.address}"
        ], 'F0FFFF')  # Lavender blush

        # Adding education section
        add_info_section(doc, 'Education', [user.education], 'F0FFFF')

        # Adding skills section
        add_info_section(doc, 'Skills', user.skills.split(','), 'F0FFFF', True)

        # Adding languages section
        add_info_section(doc, 'Languages', user.languages.split(','), 'FFF0F3', True)

        # Adding professional experience section
        add_info_section(doc, 'Professional Experience', [user.experience], 'F0FFFF')

        # Saving and sending the document
        doc_path = os.path.join(os.path.expanduser('~'), f'cv_{chat_id}.docx')
        doc.save(doc_path)
        with open(doc_path, 'rb') as doc_file:
            bot.send_document(chat_id, doc_file)
        del user_data[chat_id]  # Deleting user data after sending CV
    except Exception as e:
        bot.send_message(chat_id, f"An error occurred: {str(e)}")

# Function to set up document style
def setup_document_style(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

# Function to add user photo to the document
def add_photo(doc, photo_id):
    file_info = bot.get_file(photo_id)
    file = requests.get(f'https://api.telegram.org/file/bot{TOKEN}/{file_info.file_path}')
    image_stream = BytesIO(file.content)
    doc.add_picture(image_stream, width=Inches(2))
    # doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    last_paragraph = doc.paragraphs[-1]
    # last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    last_paragraph.style.font.size = Pt(12)  # Change font size to 12 points
    last_paragraph.style.font.color.rgb = RGBColor(0, 255, 255)  # Change font color to red (RGB: 255, 0, 0)
    last_paragraph.style.font.italic = True  # Make the font italic
    last_paragraph.style.font.underline = True  # Underline the font
    last_paragraph.style.font.name = 'Arial'  # Change font name to Arial
    last_paragraph.style.font.bold = True  # Make the font bold


# Function to add objective section to the document
def add_objective_section(doc, objective):
    objective_header = doc.add_heading('Objective', level=1)
    objective_para = doc.add_paragraph(objective)
    add_colored_background(objective_para, 'FFF0F3')

# Function to add information section to the document
def add_info_section(doc, title, content, background_color, bullet_list=False):
    section = doc.add_heading(title, level=1)
    for item in content:
        para = doc.add_paragraph()
        if bullet_list:
            para = doc.add_paragraph(item, style='ListBullet')
        else:
            para = doc.add_paragraph(item)
        add_colored_background(para, background_color)


# Function to add colored background to a paragraph
def add_colored_background(paragraph, hex_color):
    shading_el = OxmlElement('w:shd')
    shading_el.set(qn('w:fill'), hex_color)
    paragraph._p.get_or_add_pPr().append(shading_el)


# Bot command handlers and interaction logic
@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    msg = bot.reply_to(message, "Assalomu Alaykum 👋. \nWELCOME TO CV GENERATE BOT 📲\n\nENTER THE NAME: ")
    bot.register_next_step_handler(msg, process_name_step)

def process_name_step(message):
    chat_id = message.chat.id
    user_data[chat_id] = UserData()
    user_data[chat_id].name = message.text
    msg = bot.reply_to(message, "SURNAME 📝: ")
    bot.register_next_step_handler(msg, process_surname_step)

def process_surname_step(message):
    chat_id = message.chat.id
    user_data[chat_id].surname = message.text
    msg = bot.reply_to(message, "AGE 📝: ")
    bot.register_next_step_handler(msg, process_age_step)

def process_age_step(message):
    chat_id = message.chat.id
    user_data[chat_id].age = message.text
    msg = bot.reply_to(message, "ADDRESS 📝: ")
    bot.register_next_step_handler(msg, process_address_step)

def process_address_step(message):
    chat_id = message.chat.id
    user_data[chat_id].address = message.text
    msg = bot.reply_to(message, "CAREER 📝: ")
    bot.register_next_step_handler(msg, process_objective_step)

def process_objective_step(message):
    chat_id = message.chat.id
    user_data[chat_id].objective = message.text
    msg = bot.reply_to(message, "Education 📝: ")
    bot.register_next_step_handler(msg, process_education_step)

def process_education_step(message):
    chat_id = message.chat.id
    user_data[chat_id].education = message.text
    msg = bot.reply_to(message, "Specialization ❓")
    bot.register_next_step_handler(msg, process_specialization_step)

def process_specialization_step(message):
    chat_id = message.chat.id
    user_data[chat_id].specialization = message.text
    msg = bot.reply_to(message, "SKILLS 📝: =>  separated by commas")
    bot.register_next_step_handler(msg, process_skills_step)

def process_skills_step(message):
    chat_id = message.chat.id
    user_data[chat_id].skills = message.text
    msg = bot.reply_to(message, "LANGUAGES 📝: => separated by commas")
    bot.register_next_step_handler(msg, process_languages_step)

def process_languages_step(message):
    chat_id = message.chat.id
    user_data[chat_id].languages = message.text
    msg = bot.reply_to(message, "Add work experience, If you want finish click 👉 '/finish'")
    bot.register_next_step_handler(msg, process_experience_step)

def process_experience_step(message):
    chat_id = message.chat.id
    if message.text.lower() != '/finish':
        user_data[chat_id].experience += "\n" + message.text
        msg = bot.reply_to(message, "Add work experience, If you want finish click 👉 '/finish'")
        bot.register_next_step_handler(msg, process_experience_step)
    else:
        msg = bot.reply_to(message, "PLEASE SEND PHOTO 📷")
        bot.register_next_step_handler(msg, handle_photo)

def handle_photo(message):
    chat_id = message.chat.id
    if message.text == '/finish':
        create_modern_cv(chat_id)
    elif 'photo' in message.json:
        photo = message.photo[-1]
        file_id = photo.file_id
        user_data[chat_id].photo_id = file_id
        create_modern_cv(chat_id)
    else:
        msg = bot.reply_to(message, "No photo detected, please try again by uploading a photo.")
        bot.register_next_step_handler(msg, handle_photo)

# Start the bot
bot.polling(none_stop=True)
