# import telebot
# import docx
#
# # Replace with your Telegram bot token
# TOKEN = '6837439077:AAGaEwhRZCEuN6UGgIF8KKQkNG-DxemOvJg'
#
# bot = telebot.TeleBot(TOKEN)
#
# # Information storage (consider using a database for scalability)
# user_data = {}
#
# # Functions for resume creation and data manipulation
#
# def create_resume(user_id):
#     user_info = user_data.get(user_id)
#     if not user_info:
#         return "Please provide your information first."
#
#     document = docx.Document()
#
#     # Add user details (replace with actual information fields)
#     document.add_heading(user_info['name'])
#     document.add_paragraph(f"Contact: {user_info['contact']}")
#     document.add_paragraph(f"Email: {user_info['email']}")
#
#     # Add sections based on user input
#     if user_info.get('experience'):
#         document.add_heading('Experience')
#         for exp in user_info['experience']:
#             document.add_paragraph(exp)
#
#     if user_info.get('skills'):
#         document.add_heading('Skills')
#         for skill in user_info['skills']:
#             document.add_paragraph(skill)
#
#     # ... (add other sections as needed)
#
#     document.save(f"{user_id}.docx")
#     return f"Your resume is ready! Download: {user_id}.docx"  # Provide a download link (consider using a cloud storage service)
#
# def update_user_info(user_id, field, value):
#     user_data.setdefault(user_id, {})[field] = value
#
# def get_user_info(user_id, field):
#     return user_data.get(user_id, {}).get(field)
#
# # Bot interaction logic
#
# @bot.message_handler(commands=['start'])
# def start_handler(message):
#     user_id = message.chat.id
#     update_user_info(user_id, 'name', None)  # Reset user data on start
#     bot.reply_to(message, "Hi! Welcome to the Resume Builder Bot.\n"
#                            "Let's create your resume step-by-step.\n"
#                            "/name - Enter your name\n"
#                            "/contact - Enter your contact information\n"
#                            "/email - Enter your email address\n"
#                            "/experience (optional) - Add your work experience (one line per experience)\n"
#                            "/skills (optional) - Add your skills (one line per skill)\n"
#                            "/finish - Generate your resume")
#
# @bot.message_handler(commands=['name', 'contact', 'email'])
# def info_handler(message):
#     user_id = message.chat.id
#     field = message.text.split()[1]  # Extract field name from command
#     info = message.text[len(field) + 2:]  # Extract user input after command
#
#     if info:
#         update_user_info(user_id, field, info)
#         bot.reply_to(message, f"{field.capitalize()} updated!")
#     else:
#         bot.reply_to(message, f"Please enter your {field}.")
#
# @bot.message_handler(commands=['experience', 'skills'])
# def list_handler(message):
#     user_id = message.chat.id
#     field = message.text.split()[1]
#
#     if not get_user_info(user_id, field):
#         user_data[user_id][field] = []
#
#     bot.reply_to(message, f"Enter your {field} (one line per entry):")
#
# @bot.message_handler(func=lambda message: True)  # Handle all other messages
# def other_handler(message):
#     user_id = message.chat.id
#     field = get_user_info(user_id, '/experience') or get_user_info(user_id, '/skills')
#
#     if field:
#         user_data[user_id][field].append(message.text)
#         bot.reply_to(message, f"{field.capitalize()} added!")
#     else:
#         bot.reply_to(message, "I'm not sure what you mean. Use the provided commands to build your resume.")
#
# @bot.message_handler(commands=['finish'])
# def finish_handler(message):
#     user_id = message.chat.id
#
#     # Check if user provided information
#     if not get_user_info(user_id, 'name'):
#         bot.reply_to(message, "Please enter your details first using the provided commands.")
#         return
#
#     # Generate resume
#     resume_message = create_resume(user_id)
#     bot.reply_to(message, resume_message)
#
# @bot.message_handler(commands=['finish'])
# def finish_handler(message):
#     user_id = message.chat.id
#
#     # Check if user provided information
#     if not get_user_info(user_id, 'name'):
#         bot.reply_to(message, "Please enter your details first using the provided commands.")
#         return
#
#     # Generate resume (consider template file if applicable)
#     resume_message = create_resume(user_id, template_file="resume.docx")  # Uncomment if using template
#     # resume_message = create_resume(user_id)  # Use if ignoring existing file
#
#     bot.reply_to(message, resume_message)
#
#     # Consider adding download options (explained earlier)
#
# def create_resume():
#     document = docx.Document()
#
#     # Add user details
#     document.add_heading(name)
#     document.add_paragraph(f"Contact: {contact}")
#     document.add_paragraph(f"Email: {email}")
#
#     # Add experience section
#     if experience:
#         document.add_heading('Experience')
#         for exp in experience:
#             document.add_paragraph(exp)
#
#     # Add skills section
#     if skills:
#         document.add_heading('Skills')
#         for skill in skills:
#             document.add_paragraph(skill)
#
#     # ... (add other sections as needed)
#
#     document.save("predefined_resume.docx")
#     return "Your pre-defined resume is ready! Download: predefined_resume.docx"  # Consider cloud storage link
#
# # Bot interaction logic
#
# @bot.message_handler(commands=['start'])
# def start_handler(message):
#     bot.reply_to(message, "Hi! Welcome to the Pre-defined Resume Bot.\n"
#                            "This bot provides a pre-defined resume template.\n"
#                            "/download - Download the pre-defined resume")
#
# @bot.message_handler(commands=['download'])
# def download_handler(message):
#     resume_message = create_resume()
#     bot.reply_to(message, resume_message)
#
# bot.polling()




#########______RESUME__BOT____############

# import telebot
# import docx
# import io
#
# # Replace with your Telegram bot token
# TOKEN = '6837439077:AAGaEwhRZCEuN6UGgIF8KKQkNG-DxemOvJg'
#
# bot = telebot.TeleBot(TOKEN)

# # Information storage
# user_data = {}

# Function to create the resume based on user information
# def create_resume(user_id, format='docx'):
#     user_info = user_data.get(user_id)
#     if not user_info:
#         return "Please provide your information first."
#     document = docx.Document()
#     # Add user details
#     document.add_heading(user_info.get('name', 'Name not provided'), level=1)
#     document.add_paragraph(f"Contact: {user_info.get('contact', 'Contact not provided')}")
#     document.add_paragraph(f"Email: {user_info.get('email', 'Email not provided')}")
#
#     # Add experience section
#     if user_info.get('experience'):
#         document.add_heading('Experience', level=2)
#         for exp in user_info['experience']:
#             document.add_paragraph(exp)
#
#     # Add skills section
#     if user_info.get('skills'):
#         document.add_heading('Skills', level=2)
#         for skill in user_info['skills']:
#             document.add_paragraph(skill)
#
#     # ... (add other sections as needed)
#
#     if format == 'docx':
#         file_name = f"{user_id}_resume.docx"
#         document.save(file_name)
#         return file_name
#     elif format == 'pdf':
#         file_name = f"{user_id}_resume.pdf"
#         pdf_bytes = io.BytesIO()
#         document.save(pdf_bytes)
#         pdf_bytes.seek(0)
#         return file_name, pdf_bytes.getvalue()
#
#
# # Bot interaction logic
#
# @bot.message_handler(commands=['start'])
# def start_handler(message):
#     user_id = message.chat.id
#     user_data[user_id] = {}  # Initialize user data
#     bot.reply_to(message, "Hi! Welcome to the Resume Builder Bot.\n"
#                            "Let's create your resume step-by-step.\n"
#                            "Please enter your name using /name /contact /email /experience /skills command.\nYou can download click /finish")
#
# @bot.message_handler(commands=['name', 'contact', 'email', 'experience', 'skills'])
# def info_handler(message):
#     user_id = message.chat.id
#     command = message.text.split()[0][1:]  # Extract command
#     info = message.text[len(command) + 2:]  # Extract user input after command
#
#     if info:
#         user_data[user_id][command] = info
#         if command != 'skills':
#             bot.reply_to(message, f"Your {command} has been updated successfully!")
#         else:
#             bot.reply_to(message, f"Your {command[:-1]} have been updated successfully!")
#     else:
#         bot.reply_to(message, f"Please enter your {command}.")
#
# @bot.message_handler(commands=['finish'])
# def finish_handler(message):
#     user_id = message.chat.id
#
#     # Check if user provided information
#     if not all(user_data[user_id].values()):
#         bot.reply_to(message, "Please provide all your information first using the provided commands.")
#         return
#
#     # Generate resume
#     markup = telebot.types.InlineKeyboardMarkup()
#     btn_docx = telebot.types.InlineKeyboardButton(text="Download as DOCX", callback_data="download_docx")
#     btn_pdf = telebot.types.InlineKeyboardButton(text="Download as PDF", callback_data="download_pdf")
#     markup.add(btn_docx, btn_pdf)
#     bot.send_message(message.chat.id, "Choose the format for your resume:", reply_markup=markup)
#
# @bot.callback_query_handler(func=lambda call: call.data.startswith('download'))
# def download_callback_handler(call):
#     format = call.data.split('_')[-1]
#     user_id = call.message.chat.id
#     if format in ['docx', 'pdf']:
#         if format == 'pdf':
#             file_name, file_content = create_resume(user_id, format='pdf')
#             bot.send_document(user_id, file_content, filename=file_name, caption="Your resume is ready!")
#         else:
#             file_name = create_resume(user_id, format='docx')
#             with open(file_name, 'rb') as f:
#                 bot.send_document(user_id, f, caption="Your resume is ready!")
#     else:
#         bot.send_message(user_id, "Invalid format.")
#
#
# bot.polling(timeout=60)






######___________it is not working_____###########
#
# import telebot
# import docx
#
# TOKEN = '6837439077:AAGaEwhRZCEuN6UGgIF8KKQkNG-DxemOvJg'
#
# bot = telebot.TeleBot(TOKEN)
#
# user_data = {}
#
# # Function to create the resume
# def create_resume(user_id):
#     user_info = user_data.get(user_id)
#     if not user_info:
#         return "Please provide your information first."
#
#     document = docx.Document()
#
#     # Add user details
#     document.add_heading(user_info.get('name', 'Name not provided'), level=1)
#     document.add_paragraph(f"Contact: {user_info.get('contact', 'Contact not provided')}")
#     document.add_paragraph(f"Email: {user_info.get('email', 'Email not provided')}")
#     document.add_paragraph(f"Address: {user_info.get('address', 'Address not provided')}")
#
#     # Add experience section
#     if user_info.get('experience'):
#         document.add_heading('Experience', level=2)
#         for exp in user_info['experience']:
#             document.add_paragraph(exp)
#
#     # Add skills section
#     if user_info.get('skills'):
#         document.add_heading('Skills', level=2)
#         for skill in user_info['skills']:
#             document.add_paragraph(skill)
#
#     document.save(f"{user_id}_resume.docx")
#     return f"Your resume is ready! Download: {user_id}_resume.docx"
#
# # Bot interaction logic
#
# @bot.message_handler(commands=['start'])
# def start_handler(message):
#     user_id = message.chat.id
#     user_data[user_id] = {}  # Initialize user data
#     bot.reply_to(message, "Hi! Welcome to the Resume Builder Bot.\n"
#                           "Let's create your resume step-by-step.\n"
#                           "Please enter your name using /name command.")
#
# @bot.message_handler(commands=['name'])
# def name_handler(message):
#     user_id = message.chat.id
#     info = message.text.split(maxsplit=1)[1]
#     user_data[user_id]['name'] = info
#     bot.reply_to(message, "Your name has been updated successfully!\n"
#                           "Please enter your contact information using /contact command.")
#
# @bot.message_handler(commands=['contact'])
# def contact_handler(message):
#     user_id = message.chat.id
#     info = message.text.split(maxsplit=1)[1]
#     user_data[user_id]['contact'] = info
#     bot.reply_to(message, "Your contact information has been updated successfully!\n"
#                           "Please enter your email address using /email command.")
#
# @bot.message_handler(commands=['email'])
# def email_handler(message):
#     user_id = message.chat.id
#     info = message.text.split(maxsplit=1)[1]
#     user_data[user_id]['email'] = info
#     bot.reply_to(message, "Your email address has been updated successfully!\n"
#                           "Please enter your address using /address command.")
#
# @bot.message_handler(commands=['address'])
# def address_handler(message):
#     user_id = message.chat.id
#     info = message.text.split(maxsplit=1)[1]
#     user_data[user_id]['address'] = info
#     bot.reply_to(message, "Your address has been updated successfully!\n"
#                           "You can now add your work experience using /experience command.")
#
# @bot.message_handler(commands=['experience'])
# def experience_handler(message):
#     user_id = message.chat.id
#     info = message.text.split(maxsplit=1)[1]
#     user_data[user_id].setdefault('experience', []).append(info)
#     bot.reply_to(message, "Your work experience has been added!\n"
#                           "You can add more experience or move on to skills using /skills command.")
#
# @bot.message_handler(commands=['skills'])
# def skills_handler(message):
#     user_id = message.chat.id
#     info = message.text.split(maxsplit=1)[1]
#     user_data[user_id].setdefault('skills', []).append(info)
#     bot.reply_to(message, "Your skills have been added!\n"
#                           "You can add more skills or finish your resume using /finish command.")
#
# @bot.message_handler(commands=['finish'])
# def finish_handler(message):
#     user_id = message.chat.id
#     resume_message = create_resume(user_id)
#     bot.reply_to(message, resume_message)
#
# bot.polling()



# by_fotima
# import telebot
# from docx import Document
# # import emoji as emoji_lib  # Rename the import to avoid conflict
# import emoji as emoji_lib
#
# # Initialize your bot with the token
# bot = telebot.TeleBot("6837439077:AAGaEwhRZCEuN6UGgIF8KKQkNG-DxemOvJg")
#
# # Function to handle the /start command
# @bot.message_handler(commands=['start'])
# def start(message):
#     bot.reply_to(message, "Welcome to CV Generator Bot! Please send /start_cv to start creating your CV.")
#
# # Dictionary to store user data temporarily
# user_data = {}
#
# # Function to handle the /cv command
# @bot.message_handler(commands=['cv'])
# def cv(message):
#     bot.reply_to(message, 'Great! Let\'s start creating your CV. Please provide the following information:\n'
#                           '1. Name')
#
#     # Initialize the user's data
#     user_data[message.chat.id] = {}
#
# # Function to handle user messages and gather CV information
# @bot.message_handler(func=lambda message: True)
# def handle_message(message):
#     chat_id = message.chat.id
#     if chat_id not in user_data:
#         bot.send_message(chat_id, "Please start by sending /cv to create your CV.")
#         return
#
#     step = len(user_data[chat_id]) + 1
#     if step == 1:
#         user_data[chat_id]['Name'] = message.text
#         bot.send_message(chat_id, '2. Surname')
#     elif step == 2:
#         user_data[chat_id]['Surname'] = message.text
#         bot.send_message(chat_id, '3. Age')
#     elif step == 3:
#         user_data[chat_id]['Age'] = message.text
#         bot.send_message(chat_id, '4. Address')
#     elif step == 4:
#         user_data[chat_id]['Address'] = message.text
#         bot.send_message(chat_id, '5. Education')
#     elif step == 5:
#         user_data[chat_id]['Education'] = message.text
#         bot.send_message(chat_id, '6. Specialization')
#     elif step == 6:
#         user_data[chat_id]['Specialization'] = message.text
#         bot.send_message(chat_id, '7. Skills')
#     elif step == 7:
#         user_data[chat_id]['Skills'] = message.text
#         bot.send_message(chat_id, '8. Languages')
#     elif step == 8:
#         user_data[chat_id]['Languages'] = message.text
#         bot.send_message(chat_id, '9. Photo (emoji)\nPlease send the emoji you want to include in your CV.')
#     # Function to handle the photo (emoji) message
#     elif step == 9:
#         user_data[chat_id]['Photo'] = message.text
#         bot.send_message(chat_id, '10. Experience')
#     elif step == 10:
#             user_data[chat_id]['Experience'] = message.text
#             create_cv(message)
#
# # Function to create the CV
# def create_cv(message):
#     file_name = "Resume.docx"
#     try:
#         cv_doc = Document(file_name)
#     except FileNotFoundError:
#         bot.send_message(message.chat.id, "The base document 'Resume.docx' was not found.")
#         return
#
#     # Update specific fields with user input
#     for key, value in user_data[message.chat.id].items():
#         cv_doc.add_paragraph(f"{key}: {value}")
#
#     updated_file_name = f"{user_data[message.chat.id]['Name']}_{user_data[message.chat.id]['Surname']}_Updated_CV.docx"
#     cv_doc.save(updated_file_name)
#
#     with open(updated_file_name, 'rb') as doc:
#         bot.send_document(message.chat.id, doc)
#
#     user_data.pop(message.chat.id, None)
#
# # Start the bot
# bot.polling(timeout=60)









# imaage code not work
# import telebot
# from docx import Document
#
# # Initialize your bot with the token
# bot = telebot.TeleBot("6837439077:AAGaEwhRZCEuN6UGgIF8KKQkNG-DxemOvJg")
#
# # Function to handle the /start command
# @bot.message_handler(commands=['start'])
# def start(message):
#     bot.reply_to(message, "Welcome to CV Generator Bot! Please send /cv to start creating your CV.")
#
# # Dictionary to store user data temporarily
# user_data = {}
#
# # Function to handle the /cv command
# @bot.message_handler(commands=['cv'])
# def cv(message):
#     bot.reply_to(message, 'Great! Let\'s start creating your CV. Please provide the following information:\n'
#                           '1. Name')
#
#     # Initialize the user's data
#     user_data[message.chat.id] = {}
#
# # Function to handle user messages and gather CV information
# @bot.message_handler(func=lambda message: True)
# def handle_message(message):
#     chat_id = message.chat.id
#     if chat_id not in user_data:
#         bot.send_message(chat_id, "Please start by sending /cv to create your CV.")
#         return
#
#     step = len(user_data[chat_id]) + 1
#     if step == 1:
#         user_data[chat_id]['Name'] = message.text
#         bot.send_message(chat_id, '2. Surname')
#     elif step == 2:
#         user_data[chat_id]['Surname'] = message.text
#         bot.send_message(chat_id, '3. Age')
#     elif step == 3:
#         user_data[chat_id]['Age'] = message.text
#         bot.send_message(chat_id, '4. Address')
#     elif step == 4:
#         user_data[chat_id]['Address'] = message.text
#         bot.send_message(chat_id, '5. Education')
#     elif step == 5:
#         user_data[chat_id]['Education'] = message.text
#         bot.send_message(chat_id, '6. Specialization')
#     elif step == 6:
#         user_data[chat_id]['Specialization'] = message.text
#         bot.send_message(chat_id, '7. Skills')
#     elif step == 7:
#         user_data[chat_id]['Skills'] = message.text
#         bot.send_message(chat_id, '8. Languages')
#     elif step == 8:
#         user_data[chat_id]['Languages'] = message.text
#         bot.send_message(chat_id, '9. Photo (emoji)\nPlease send the emoji you want to include in your CV.')
#     elif step == 9:
#         user_data[chat_id]['Photo'] = message.text
#         bot.send_message(chat_id, '10. Experience')
#
#     elif step == 10:
#         user_data[chat_id]['Experience'] = message.text
#         create_cv(message)
#
# # Function to create the CV
# def create_cv(message):
#     file_name = "Resume.docx"
#     try:
#         cv_doc = Document(file_name)
#     except FileNotFoundError:
#         bot.send_message(message.chat.id, "The base document 'Resume.docx' was not found.")
#         return
#
#     # Update specific fields with user input
#     for key, value in user_data[message.chat.id].items():
#         cv_doc.add_paragraph(f"{key}: {value}")
#
#     updated_file_name = f"{user_data[message.chat.id]['Name']}_{user_data[message.chat.id]['Surname']}_Updated_CV.docx"
#     cv_doc.save(updated_file_name)
#
#     with open(updated_file_name, 'rb') as doc:
#         bot.send_document(message.chat.id, doc)
#
#     user_data.pop(message.chat.id, None)
#
# # Start the bot
# bot.polling(timeout=60)

import telebot
from docx import Document
from io import BytesIO
from docx.shared import Pt
from docx.shared import RGBColor

# Initialize your bot with the token
bot = telebot.TeleBot("6837439077:AAGaEwhRZCEuN6UGgIF8KKQkNG-DxemOvJg")


# Function to handle the /start command
@bot.message_handler(commands=['start'])
def start(message):
    bot.reply_to(message, "Welcome to CV Generator Bot! Please send /cv to start creating your CV.")

# Dictionary to store user data temporarily
user_data = {}

# Function to handle the /cv command
@bot.message_handler(commands=['cv'])
def cv(message):
    bot.reply_to(message, 'Great! Let\'s start creating your CV. Please provide the following information:\n'
                          '1. Name')

    # Initialize the user's data
    user_data[message.chat.id] = {}

# Function to handle user messages and gather CV information
@bot.message_handler(func=lambda message: True)
def handle_message(message):
    chat_id = message.chat.id
    if chat_id not in user_data:
        bot.send_message(chat_id, "Please start by sending /cv to create your CV.")
        return

    step = len(user_data[chat_id]) + 1
    if step == 1:
        user_data[chat_id]['Name'] = message.text
        bot.send_message(chat_id, '2. Surname')
    elif step == 2:
        user_data[chat_id]['Surname'] = message.text
        bot.send_message(chat_id, '3. Age')
    elif step == 3:
        user_data[chat_id]['Age'] = message.text
        bot.send_message(chat_id, '4. Address')
    elif step == 4:
        user_data[chat_id]['Address'] = message.text
        bot.send_message(chat_id, '5. Education')
    elif step == 5:
        user_data[chat_id]['Education'] = message.text
        bot.send_message(chat_id, '6. Specialization')
    elif step == 6:
        user_data[chat_id]['Specialization'] = message.text
        bot.send_message(chat_id, '7. Skills')
    elif step == 7:
        user_data[chat_id]['Skills'] = message.text
        bot.send_message(chat_id, '8. Languages')
    elif step == 8:
        user_data[chat_id]['Languages'] = message.text
        bot.send_message(chat_id, '9. Experience')
    elif step == 9:
        user_data[chat_id]['Experience'] = message.text
        bot.send_message(chat_id, '10. Please send the photo you want to include in your CV.')

# Function to handle photo messages and add them to the CV
@bot.message_handler(content_types=['photo'])
def handle_photo(message):
    chat_id = message.chat.id
    if chat_id not in user_data:
        bot.send_message(chat_id, "Please start by sending /cv to create your CV.")
        return

    step = len(user_data[chat_id]) + 1
    if step == 10:
        photo_info = message.photo[-1]  # Get the largest photo
        photo_data = bot.download_file(bot.get_file(photo_info.file_id).file_path)
        user_data[chat_id]['Photo'] = photo_data
        bot.send_message(chat_id, 'Your CV is being generated...')
        create_cv(message)
    else:
        bot.send_message(chat_id, 'Please send the photo at the appropriate step.')

# Function to create the CV
def create_cv(message):
    file_name = "Resume.docx"
    try:
        cv_doc = Document(file_name)
    except FileNotFoundError:
        bot.send_message(message.chat.id, "The base document 'Resume.docx' was not found.")
        return

    # Update specific fields with user input
    for key, value in user_data[message.chat.id].items():
        if key == 'Photo':
            photo_stream = BytesIO(value)
            cv_doc.add_picture(photo_stream, width=200, height=200)
        else:
            p = cv_doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True  # Set text to bold
            p.add_run(value).italic = True  # Set text to italic
            p.add_run('\n')  # Add newline for spacing
            # Set font size to 18pt for the value (last run)
            for run in p.runs:
                run.font.size = Pt(18)
                break  # Apply the font size only to the value, then exit the loop
            # Set font size to 20pt for the key (second-last run)
            p.runs[-2].font.size = Pt(20)
            # Set font color to blue for the value (last run)
            p.runs[-1].font.color.rgb = RGBColor(0, 0, 255)  # RGB color code for blue

    updated_file_name = f"{user_data[message.chat.id]['Name']}_{user_data[message.chat.id]['Surname']}_Updated_CV.docx"
    cv_doc.save(updated_file_name)

    with open(updated_file_name, 'rb') as doc:
        bot.send_document(message.chat.id, doc)

    user_data.pop(message.chat.id, None)

# Start the bot
bot.polling(timeout=60)
