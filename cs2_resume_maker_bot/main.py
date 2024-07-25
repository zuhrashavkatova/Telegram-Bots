# import docx
# import datetime
#
# my_doc = docx.Document("resume.docx")
# all_paras = my_doc.paragraphs
#
# for x in (all_paras):
#     print(x.test)
#
# # only by number of paragraph
# specific_para = my_doc_paragraph[-1]
# print(type(specific_para))
# for run in specific_para.runs:
#     print(run.text)
# from xml.dom.minidom import Document
#
import docx

# my_doc = docx.Document("resume.docx")
# def change_table_values(doc_path, new_text, row_idx, col_idx, save_as):
#     doc = Document(doc_path)
#     table = doc.table[0]
#     call = table.cell(row_idx, col_idx)
#     call.text = new_text
# def save_doc(st_name, teacger_name, learning_aims, date_sub, comment, deadline):
#     change_table_values("resume.docx",st_name,1,4, f"{}.docx")
#     change_table_values(f"{}.docx", teacger_name, 2, 4, f"{}.docx")
#     change_table_values()



# Read document
my_doc = docx.Document('resume.docx')
# all_paras = my_doc.paragraphs
# # print(len(all_paras))
# for x in (all_paras):
#     print(x.text)

# specific_para = my_doc.paragraphs[-1]
# print(type(specific_para))
# for run in specific_para.runs:
#     print(run.text)


# my_doc.add_paragraph("IMPACT 205 is high")
# specific_para = ""
# var1 = specific_para.add_run("NEW TEXT")
# var1.font.name = "Impact"
# my_doc.save("resume.docx")






import docx
import datetime
import Doc
my_doc = docx.Document('resume.docx')
my_doc = docx.Document("HI MY NAME.docx")
all_paras = my_doc.paragraphs
print(len(all_paras))

specific_para = my_doc.paragraphs[-2]
# print(specific_para.text)
print(type(specific_para))
for run in specific_para.runs:
    print(run.text)



my_doc.add_paragraph("IMPACT 205 is high ranked failures")
specific_para = ""

var1 = specific_para.add_run("Harry")
var1.font.name = "Harry P"
my_doc.save("White-And-Brown-Vintage-Resume-_1_.docx")



p2 = my_doc.add_paragraph('')
run2 = p2.add_run('this is Harry Potter')
run2.font.name = 'Harry P'
my_doc.save("White-And-Brown-Vintage-Resume-_1_.docx")

all_paras = my_doc.paragraphs
for x in (all_paras):
    print(x.text)

document = Document()
paragraph = document.add_pragraph('Lorem ipsum dolor sit amet.')
document.add_heading('the REAL meaning of the universe')
document.add_heading('The role of dolphins', level=1)
document.save("White-And-Brown-Vintage-Resume-_1_.docx")

for paragraph in document.paragraphs:
    if 'dolphins' in paragraph.text:
        paragraph.text = paragraph.text.replace('dolphins', 'Rabbit')

document.save('White-And-Brown-Vintage-Resume-_1_.docx')

def save_doc(st_name, teacher_name, learning_aims, date_sub, comment, deadline):
    change_table_values("White And Brown Vintage Resume (1).pdf", st_name, 1, 4, f'{st_name}.docx')
    change_table_values(f"{st_name}.docx", teacher_name, 1, 4, f'{st_name}.docx')
    change_table_values(f"{st_name}.docx", learning_aims, 1, 4, f'{st_name}.docx')
    change_table_values(f"{st_name}.docx", date_sub, 1, 4, f'{st_name}.docx')
    change_table_values(f"{st_name}.docx", deadline, 1, 4, f'{st_name}.docx')
    change_table_values(f"{st_name}.docx", comment, 1, 4, f'{st_name}.docx')
    change_table_values(f"{st_name}.docx", st_name, 1, 4, f'{st_name}.docx')
    change_table_values(f"{st_name}.docx", date_sub, 1, 4, f'{st_name}.docx')
    change_table_values(f"{st_name}.docx", teacher_name, 1, 4, f'{st_name}.docx')
    change_table_values(f"{st_name}.docx", teacher_name, 1, 4, f'{st_name}.docx')
    change_table_values(f"{st_name}.docx", my_date, 1, 4, f'{st_name}.docx')
    change_table_values(f"{st_name}.docx", my_date, 1, 4, f'{st_name}.docx')


import csv
import datetime
x = datetime.datetime.now()
my_date = x.strftime("%d/%B/%y")
with open("White-And-Brown-Vintage-Resume-_1_.docx", "r", newline="") as file:
 csv_reader = csv.reader(file)
 for x in csv_reader:
    if x[0] == "id":
        continue
    save_doc(f"{x[1]} {x[2]}","Fotima Shavkatova", "L01 L02", x[4], f"Grade {x[3]} \n")



